from flask import Flask, request, render_template, session, send_file, redirect, url_for
import google.generativeai as genai
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

app = Flask(__name__)
app.secret_key = "a9f4b0d2c7e8f13d6b4c1a9d8e7f3a1b"
app.config["UPLOAD_FOLDER"] = "uploads"

# Direct API key (⚠ replace with your own)
API_KEY = "AIzaSyBywhW6RSkhvswm9yN2Mht8x5jbK-6pqYM"
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")

# Ensure uploads folder exists
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

def save_to_txt(text):
    """Save plain text history."""
    with open("history.txt", "a", encoding="utf-8") as f:
        f.write(text + "\n")

# --- PDF Text Extraction ---
def extract_pdf_text(file_path):
    """Extract text from a PDF, using OCR if needed."""
    text = ""
    pdf_document = fitz.open(file_path)
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        page_text = page.get_text()

        if not page_text.strip():
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img)

        text += page_text + "\n"
    return text

# --- Main Translation Route ---
@app.route("/", methods=["GET", "POST"])
def translate_text():
    if request.method == "POST":
        target_language = request.form.get("target_lang", "Urdu")
        user_input = ""
        translated_text = ""

        # Text input translation
        if "message" in request.form and request.form["message"].strip():
            user_input = request.form["message"]
            prompt = f"Translate the following text to {target_language}:\n\n{user_input}"
            response = model.generate_content(prompt)
            translated_text = response.text

            save_to_txt(user_input)

        # PDF file translation
        elif "pdf_file" in request.files:
            pdf_file = request.files["pdf_file"]
            if pdf_file.filename != "":
                file_path = os.path.join(app.config["UPLOAD_FOLDER"], pdf_file.filename)
                pdf_file.save(file_path)

                pdf_text = extract_pdf_text(file_path)
                if not pdf_text.strip():
                    translated_text = "Could not extract text from the PDF."
                else:
                    user_input = pdf_text
                    prompt = f"Translate the following PDF content to {target_language}:\n\n{pdf_text[:10000]}"
                    response = model.generate_content(prompt)
                    translated_text = response.text

                    save_to_txt(pdf_text)

        session["user_input"] = user_input
        session["translated_text"] = translated_text
        session["target_language"] = target_language

    return render_template(
        "index.html",
        user_input=session.get("user_input", ""),
        translated_text=session.get("translated_text", ""),
        target_language=session.get("target_language", "Urdu")
    )

# ---------- File Conversion Routes ----------
@app.route("/pdf-to-docx", methods=["POST"])
def pdf_to_docx():
    pdf_file = request.files["pdf_file"]
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], pdf_file.filename)
    pdf_file.save(file_path)

    pdf_text = extract_pdf_text(file_path)
    doc = Document()
    doc.add_paragraph(pdf_text)
    docx_path = file_path.replace(".pdf", ".docx")
    doc.save(docx_path)

    return send_file(docx_path, as_attachment=True)

@app.route("/docx-to-pdf", methods=["POST"])
def docx_to_pdf():
    docx_file = request.files["docx_file"]
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], docx_file.filename)
    docx_file.save(file_path)

    doc = Document(file_path)
    pdf_path = file_path.replace(".docx", ".pdf")

    c = canvas.Canvas(pdf_path, pagesize=letter)
    text_obj = c.beginText(40, 750)
    text_obj.setFont("Times-Roman", 12)

    for para in doc.paragraphs:
        text_obj.textLine(para.text)

    c.drawText(text_obj)
    c.save()

    return send_file(pdf_path, as_attachment=True)

@app.route("/image-to-pdf", methods=["POST"])
def image_to_pdf():
    image_file = request.files["image_file"]
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], image_file.filename)
    image_file.save(file_path)

    img = Image.open(file_path)
    pdf_path = file_path.rsplit(".", 1)[0] + ".pdf"
    img.convert("RGB").save(pdf_path)

    return send_file(pdf_path, as_attachment=True)

# ---------- Show Translation History ----------
# @app.route("/show-history", methods=["GET"])
# def show_history():
#     try:
#         with open("history.txt", "r", encoding="utf-8") as f:
#             history = f.read()
        
#         if not history.strip():
#             return """
#             <html>
#             <body style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
#                 <h2>No Translation History</h2>
#                 <p>No translations have been saved yet.</p>
#                 <a href="/" style="background-color: #4CAF50; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">← Back to Home</a>
#             </body>
#             </html>
#             """
        
#         entries = history.strip().split('\n')
        
#         html_content = f"""
#         <html>
#         <head>
#             <title>Translation History - GeoSpeak</title>
#             <style>
#                 body {{ 
#                     font-family: Arial, sans-serif; 
#                     margin: 20px; 
#                     background-color: #f5f5f5;
#                 }}
#                 .container {{
#                     max-width: 800px;
#                     margin: 0 auto;
#                     background-color: white;
#                     padding: 20px;
#                     border-radius: 10px;
#                     box-shadow: 0 2px 10px rgba(0,0,0,0.1);
#                 }}
#                 .entry {{ 
#                     margin-bottom: 15px; 
#                     border: 1px solid #ddd; 
#                     padding: 15px; 
#                     border-radius: 8px;
#                     background-color: #fafafa;
#                 }}
#                 h1 {{ 
#                     color: #4CAF50; 
#                     text-align: center;
#                     margin-bottom: 30px;
#                 }}
#                 .stats {{
#                     background-color: #e8f5e8;
#                     padding: 10px;
#                     border-radius: 5px;
#                     margin-bottom: 20px;
#                     text-align: center;
#                 }}
#                 .back-btn {{
#                     display: inline-block;
#                     background-color: #4CAF50;
#                     color: white;
#                     padding: 10px 20px;
#                     text-decoration: none;
#                     border-radius: 5px;
#                     margin-top: 20px;
#                 }}
#                 .back-btn:hover {{
#                     background-color: #45a049;
#                 }}
#             </style>
#         </head>
#         <body>
#             <div class="container">
#                 <h1>📋 Translation History</h1>
#                 <div class="stats">
#                     <strong>Total Entries:</strong> {len(entries)}
#                 </div>
#         """
        
#         for i, entry in enumerate(reversed(entries), 1):
#             if entry.strip():
#                 display_text = entry[:200] + "..." if len(entry) > 200 else entry
#                 html_content += f"""
#                 <div class="entry">
#                     <h3>📝 Entry #{i}</h3>
#                     <p>{display_text}</p>
#                 </div>
#                 """
        
#         html_content += """
#                 <a href="/" class="back-btn">← Back to Home</a>
#             </div>
#         </body>
#         </html>
#         """
        
#         return html_content
        
#     except FileNotFoundError:
#         return """
#         <html>
#         <body style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
#             <h2>No History File Found</h2>
#             <p>Start translating to create history!</p>
#             <a href="/" style="background-color: #4CAF50; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">← Back to Home</a>
#         </body>
#         </html>
#         """

if __name__ == "__main__":
    app.run(debug=True)

# from flask import Flask, request, render_template, session, send_file, redirect, url_for
# import google.generativeai as genai
# import fitz  # PyMuPDF
# import pytesseract
# from PIL import Image
# import os
# from docx import Document
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter

# # --- Vector DB Imports ---
# import faiss
# import numpy as np
# from sentence_transformers import SentenceTransformer

# app = Flask(__name__)
# app.secret_key = "a9f4b0d2c7e8f13d6b4c1a9d8e7f3a1b"
# app.config["UPLOAD_FOLDER"] = "uploads"

# # Direct API key (⚠ replace with your own)
# API_KEY = "AIzaSyBywhW6RSkhvswm9yN2Mht8x5jbK-6pqYM"
# genai.configure(api_key=API_KEY)
# model = genai.GenerativeModel("gemini-2.5-flash")

# # Ensure uploads folder exists
# os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# # --- Setup FAISS Vector DB ---
# embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
# dimension = 384  # embedding size for this model

# if not os.path.exists("vector.index"):
#     index = faiss.IndexFlatL2(dimension)
#     faiss.write_index(index, "vector.index")
# else:
#     index = faiss.read_index("vector.index")

# def save_to_vector_db(text):
#     """Save text embedding to FAISS and also to a text file."""
#     emb = embedding_model.encode([text])
#     index.add(np.array(emb, dtype="float32"))
#     faiss.write_index(index, "vector.index")

#     # Save vector + text to TXT file
#     with open("vector_history.txt", "a", encoding="utf-8") as vf:
#         vf.write(f"Text: {text}\nVector: {emb.tolist()}\n\n")

#     # Debug print
#     print(f"\n[Vector Saved] Text: {text[:50]}...")
#     print("Vector:", emb)

# def save_to_txt(text):
#     """Save plain text history."""
#     with open("history.txt", "a", encoding="utf-8") as f:
#         f.write(text + "\n")

# # --- PDF Text Extraction ---
# def extract_pdf_text(file_path):
#     """Extract text from a PDF, using OCR if needed."""
#     text = ""
#     pdf_document = fitz.open(file_path)
#     for page_num in range(len(pdf_document)):
#         page = pdf_document[page_num]
#         page_text = page.get_text()

#         if not page_text.strip():
#             pix = page.get_pixmap()
#             img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#             page_text = pytesseract.image_to_string(img)

#         text += page_text + "\n"
#     return text
# # @app.route("/", methods=["GET", "POST"])
# # def index():
# #     # if not model_loaded:
# #     #     return render_template_string(f"""
# #     #         <html>
# #     #             <head>
# #     #                 <meta http-equiv="refresh" content="1">
# #     #                 <title>Loading...</title>
# #     #                 <style>
# #     #                     body {{
# #     #                         font-family: Arial, sans-serif;
# #     #                         text-align: center;
# #     #                         background-color: #f7f7f7;
# #     #                         padding-top: 10%;
# #     #                     }}
# #     #                     img {{
# #     #                         width: 150px;
# #     #                         margin-bottom: 20px;
# #     #                     }}
# #     #                     h1 {{
# #     #                         color: #333;
# #     #                     }}
# #     #                     p {{
# #     #                         color: #555;
# #     #                     }}
# #     #                 </style>
# #     #             </head>
# #     #             <body>
# #     #                 <img src="{url_for('static', filename='geospeak.png')}" alt="GeoSpeak Logo">
# #     #                 <h1>🚀 Loading GeoSpeak...</h1>
# #     #                 <p>Please wait while we prepare your application.</p>
# #     #             </body>
# #     #         </html>
# #     #     """)

# #     if request.method == "POST":
# #         text = request.form.get("text")
# #         if text:
# #             save_to_vector_db(text)

# #             # Use Gemini for a quick response
# #             response = genai.GenerativeModel("gemini-pro").generate_content(
# #                 f"Summarize this text: {text}"
# #             )

# #             return render_template_string("""
# #                 <h1>✅ Data Saved & Summarized</h1>
# #                 <p><b>Original:</b> {{orig}}</p>
# #                 <p><b>Summary:</b> {{summ}}</p>
# #                 <a href="/">Go Back</a>
# #             """, orig=text, summ=response.text)

# #     return render_template_string("""
# #         <h1>FAISS + Gemini App</h1>
# #         <form method="POST">
# #             <textarea name="text" rows="5" cols="40" placeholder="Enter text here"></textarea><br>
# #             <button type="submit">Save & Summarize</button>
# #         </form>
# #     """)

# # --- Main Translation Route ---
# @app.route("/", methods=["GET", "POST"])
# def translate_text():
#     if request.method == "POST":
#         target_language = request.form.get("target_lang", "Urdu")
#         user_input = ""
#         translated_text = ""

#         # Text input translation
#         if "message" in request.form and request.form["message"].strip():
#             user_input = request.form["message"]
#             prompt = f"Translate the following text to {target_language}:\n\n{user_input}"
#             response = model.generate_content(prompt)
#             translated_text = response.text

#             save_to_vector_db(user_input)
#             save_to_txt(user_input)

#         # PDF file translation
#         elif "pdf_file" in request.files:
#             pdf_file = request.files["pdf_file"]
#             if pdf_file.filename != "":
#                 file_path = os.path.join(app.config["UPLOAD_FOLDER"], pdf_file.filename)
#                 pdf_file.save(file_path)

#                 pdf_text = extract_pdf_text(file_path)
#                 if not pdf_text.strip():
#                     translated_text = "Could not extract text from the PDF."
#                 else:
#                     user_input = pdf_text
#                     prompt = f"Translate the following PDF content to {target_language}:\n\n{pdf_text[:10000]}"
#                     response = model.generate_content(prompt)
#                     translated_text = response.text

#                     save_to_vector_db(pdf_text)
#                     save_to_txt(pdf_text)

#         session["user_input"] = user_input
#         session["translated_text"] = translated_text
#         session["target_language"] = target_language

#     return render_template(
#         "index.html",
#         user_input=session.get("user_input", ""),
#         translated_text=session.get("translated_text", ""),
#         target_language=session.get("target_language", "Urdu")
#     )

# # ---------- File Conversion Routes ----------
# @app.route("/pdf-to-docx", methods=["POST"])
# def pdf_to_docx():
#     pdf_file = request.files["pdf_file"]
#     file_path = os.path.join(app.config["UPLOAD_FOLDER"], pdf_file.filename)
#     pdf_file.save(file_path)

#     pdf_text = extract_pdf_text(file_path)
#     doc = Document()
#     doc.add_paragraph(pdf_text)
#     docx_path = file_path.replace(".pdf", ".docx")
#     doc.save(docx_path)

#     return send_file(docx_path, as_attachment=True)

# @app.route("/docx-to-pdf", methods=["POST"])
# def docx_to_pdf():
#     docx_file = request.files["docx_file"]
#     file_path = os.path.join(app.config["UPLOAD_FOLDER"], docx_file.filename)
#     docx_file.save(file_path)

#     doc = Document(file_path)
#     pdf_path = file_path.replace(".docx", ".pdf")

#     c = canvas.Canvas(pdf_path, pagesize=letter)
#     text_obj = c.beginText(40, 750)
#     text_obj.setFont("Times-Roman", 12)

#     for para in doc.paragraphs:
#         text_obj.textLine(para.text)

#     c.drawText(text_obj)
#     c.save()

#     return send_file(pdf_path, as_attachment=True)

# @app.route("/image-to-pdf", methods=["POST"])
# def image_to_pdf():
#     image_file = request.files["image_file"]
#     file_path = os.path.join(app.config["UPLOAD_FOLDER"], image_file.filename)
#     image_file.save(file_path)

#     img = Image.open(file_path)
#     pdf_path = file_path.rsplit(".", 1)[0] + ".pdf"
#     img.convert("RGB").save(pdf_path)

#     return send_file(pdf_path, as_attachment=True)

# # ---------- Show Stored Vectors ----------
# # @app.route("/show-vectors", methods=["GET"])
# # def show_vectors():
# #     if index.ntotal == 0:
# #         return "No vectors stored yet."

# #     vectors = index.reconstruct_n(0, index.ntotal)  # retrieve all vectors
# #     return f"<pre>{vectors}</pre>"

# # @app.route("/splash")
# # def splash():
# #     session["splash_seen"] = True
# #     return render_template("splash.html")

# # @app.before_request
# # def show_splash_once():
# #     if request.endpoint == "static":
# #         return  # Don't redirect static files
# #     if not session.get("splash_seen") and request.endpoint != "splash":
# #         return redirect(url_for("splash"))

# if __name__ == "__main__":
#     app.run(debug=True)
